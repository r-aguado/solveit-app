"""
Genera el TFC en formato Word (.docx) siguiendo la plantilla oficial
"Modelo TFC TECNOLOGÍA 24-25" con:
- Logo "La Otra FP | PRO2" en esquina superior derecha de cada página
- Cuadraditos decorativos en esquina superior izquierda de cada página
- Numeración de página en el pie
- Times New Roman 12 / interlineado 1.5 / márgenes 2/2/2.5/2.5
- Estructura: Portada, Resumen, Índice, secciones 1-9
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\r.aguado\TFC_Rubén_Aguado\tfc"
ASSETS = os.path.join(BASE, "assets")
OUTPUT_DOCX = os.path.join(BASE, "TFC_SolveIT.docx")

# ─────────────────────────────────────────────────────────────────────────────
# Documento + márgenes
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.header_distance = Cm(0.8)
section.footer_distance = Cm(0.8)

# Estilo Normal por defecto
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.space_after = Pt(6)


# ─────────────────────────────────────────────────────────────────────────────
# HEADER en todas las páginas: cuadraditos a la izquierda + logo a la derecha
# ─────────────────────────────────────────────────────────────────────────────
def setup_header():
    header = section.header
    # Una tabla 1x2 en el header para alinear izq/der
    table = header.add_table(rows=1, cols=2, width=Cm(16))
    table.autofit = False
    table.allow_autofit = False

    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)

    # Anchos de columna
    cell_left.width = Cm(8)
    cell_right.width = Cm(8)

    # Imagen izquierda (cuadraditos)
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left = p_left.add_run()
    run_left.add_picture(os.path.join(ASSETS, "logo_corner.png"), width=Cm(2))

    # Imagen derecha (logo)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_right = p_right.add_run()
    run_right.add_picture(os.path.join(ASSETS, "logo_header.png"), width=Cm(3.5))

    # Sin bordes en la tabla del header
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def setup_footer_page_number():
    """Número de página en el pie centrado."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    # Campo PAGE
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


setup_header()
setup_footer_page_number()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def add_h1(text, page_break_before=True):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True
    return p


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.bold = True
    return p


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    run.italic = True
    return p


def add_para(text, justify=True, bold=False, italic=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────────────────────────────────────
# PORTADA (siguiendo el modelo oficial)
# ─────────────────────────────────────────────────────────────────────────────
for _ in range(3):
    doc.add_paragraph()

# Logo grande centrado
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture(os.path.join(ASSETS, "logo_portada.png"), width=Cm(8))

doc.add_paragraph()

# TRABAJO FIN DE CICLO
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TRABAJO FIN DE CICLO")
run.font.name = "Times New Roman"
run.font.size = Pt(20)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CURSO 2024 / 2025")
run.font.name = "Times New Roman"
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

# Título del trabajo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "SolveIT: Plataforma SaaS multi-tenant para la gestión inteligente "
    "de incidencias IT en pequeñas y medianas empresas"
)
run.font.name = "Times New Roman"
run.font.size = Pt(15)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Alumno/a
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Alumno/a:")
run.font.name = "Times New Roman"
run.font.size = Pt(13)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RUBÉN AGUADO PERULERO")
run.font.name = "Times New Roman"
run.font.size = Pt(13)
run.bold = True

doc.add_paragraph()

# Tutor/a
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tutor/a:")
run.font.name = "Times New Roman"
run.font.size = Pt(13)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[NOMBRE Y APELLIDOS DEL TUTOR]")
run.font.name = "Times New Roman"
run.font.size = Pt(13)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Ciclo formativo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "CICLO FORMATIVO DE GRADO SUPERIOR EN DESARROLLO DE APLICACIONES MULTIPLATAFORMA (DAM)"
)
run.font.name = "Times New Roman"
run.font.size = Pt(13)
run.bold = True

# ─────────────────────────────────────────────────────────────────────────────
# PÁGINA EN BLANCO (como en el modelo)
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# RESUMEN
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Resumen")
run.font.name = "Times New Roman"
run.font.size = Pt(14)
run.bold = True

add_para(
    "El presente Trabajo Fin de Ciclo desarrolla SolveIT, una plataforma SaaS multi-tenant "
    "dirigida a pequeñas y medianas empresas que necesitan gestionar incidencias informáticas "
    "de forma profesional sin asumir el coste de soluciones empresariales como ServiceNow o "
    "Jira Service Management. La solución integra una aplicación móvil multiplataforma "
    "desarrollada en React Native, una aplicación web responsive en Next.js y un backend "
    "Node.js con base de datos PostgreSQL. Entre sus funcionalidades destacan la asignación "
    "automática de técnicos según carga de trabajo, una base de conocimiento colaborativa, "
    "un foro interno para los equipos, recuperación segura de contraseñas mediante código "
    "por correo electrónico y, como diferenciador clave, un asistente de inteligencia "
    "artificial integrado mediante el modelo Gemma de Google que orienta a los usuarios "
    "ante problemas técnicos y reduce la carga del departamento de soporte. La arquitectura "
    "multi-tenant permite que un único despliegue dé servicio a múltiples empresas con "
    "aislamiento total de datos. La plataforma se aloja en infraestructura cloud Railway "
    "con escalado bajo demanda y un coste operativo reducido. Tras desarrollar y desplegar "
    "la plataforma, se ha validado su funcionalidad sobre dos empresas piloto, demostrando "
    "que la solución es viable como producto comercial accesible para el tejido empresarial español."
)

p = doc.add_paragraph()
run = p.add_run("Palabras clave: ")
run.font.name = "Times New Roman"
run.font.size = Pt(12)
run.bold = True
run = p.add_run("asistente IA; gestión de incidencias; multi-tenant; PYME; SaaS")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Abstract")
run.font.name = "Times New Roman"
run.font.size = Pt(14)
run.bold = True
run.italic = True

add_para(
    "This Final Project develops SolveIT, a multi-tenant SaaS platform aimed at small and "
    "medium-sized enterprises that need to manage IT incidents professionally without "
    "assuming the cost of enterprise solutions such as ServiceNow or Jira Service Management. "
    "The solution integrates a cross-platform mobile application developed in React Native, "
    "a responsive web application built in Next.js and a Node.js backend with a PostgreSQL "
    "database. Its features include automatic technician assignment based on workload, a "
    "collaborative knowledge base, an internal forum for teams, secure password recovery "
    "through email codes and, as a key differentiator, an integrated artificial intelligence "
    "assistant powered by Google's Gemma model that guides users through technical issues "
    "and reduces the workload of the support department. The multi-tenant architecture "
    "enables a single deployment to serve multiple companies with full data isolation. The "
    "platform is hosted on Railway cloud infrastructure with on-demand scaling and a reduced "
    "operating cost. After developing and deploying the platform, its functionality has "
    "been validated on two pilot companies, demonstrating that the solution is viable as a "
    "commercial product accessible to the Spanish business sector."
)

p = doc.add_paragraph()
run = p.add_run("Key words: ")
run.font.name = "Times New Roman"
run.font.size = Pt(12)
run.bold = True
run.italic = True
run = p.add_run("AI assistant; incident management; multi-tenant; SaaS; SME")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# ÍNDICE
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ÍNDICE")
run.font.name = "Times New Roman"
run.font.size = Pt(20)
run.bold = True

doc.add_paragraph()

# Tabla de contenido manual con estilo
indice_items = [
    ("1.", "INTRODUCCIÓN", "5"),
    ("2.", "ANÁLISIS", "6"),
    ("", "    2.1 Estudio de mercado", "6"),
    ("", "    2.2 Elección de soluciones", "8"),
    ("3.", "VALORACIÓN ECONÓMICA", "10"),
    ("", "    3.1 Hardware", "10"),
    ("", "    3.2 Licencias y servicios", "11"),
    ("", "    3.3 Personal", "11"),
    ("", "    3.4 Costes de mantenimiento", "12"),
    ("4.", "VIABILIDAD DEL PROYECTO", "13"),
    ("5.", "PUESTA EN MARCHA", "16"),
    ("", "    5.1 Forma jurídica y trámites", "16"),
    ("", "    5.2 Recursos humanos", "17"),
    ("", "    5.3 Prevención de riesgos laborales", "18"),
    ("", "    5.4 Análisis medioambiental", "19"),
    ("", "    5.5 Plan de financiación", "19"),
    ("6.", "EVALUACIÓN FUNCIONAL", "21"),
    ("7.", "CONCLUSIÓN", "24"),
    ("8.", "BIBLIOGRAFÍA", "25"),
    ("9.", "ANEXOS", "27"),
]

for num, titulo, pag in indice_items:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
    is_main = num != ""
    txt = f"{num} {titulo}" if is_main else titulo
    run = p.add_run(txt)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if is_main:
        run.bold = True
    p.add_run("\t").font.size = Pt(12)
    run_pag = p.add_run(pag)
    run_pag.font.name = "Times New Roman"
    run_pag.font.size = Pt(12)
    if is_main:
        run_pag.bold = True

# ─────────────────────────────────────────────────────────────────────────────
# 1. INTRODUCCIÓN
# ─────────────────────────────────────────────────────────────────────────────
add_h1("1. INTRODUCCIÓN")

add_para(
    "La transformación digital del tejido empresarial español ha incrementado de forma "
    "exponencial la dependencia de los sistemas informáticos en cualquier organización, "
    "independientemente de su tamaño o sector. Según el Instituto Nacional de Estadística "
    "(INE, 2024), más del 99% de las empresas españolas con asalariados dispone de conexión "
    "a internet y utiliza ordenadores en su actividad diaria. Esta dependencia tecnológica "
    "ha generado, paralelamente, una necesidad creciente de gestionar de forma eficiente "
    "las incidencias técnicas que inevitablemente surgen: equipos averiados, problemas de "
    "red, software que falla, accesos denegados, dispositivos periféricos que dejan de "
    "funcionar y un largo etcétera."
)

add_para(
    "Mientras que las grandes corporaciones cuentan con soluciones consolidadas como "
    "ServiceNow, BMC Helix o Jira Service Management para gestionar sus tickets de soporte, "
    "las pequeñas y medianas empresas (PYME), que representan el 99,8% del tejido "
    "empresarial español según el Ministerio de Industria (2024), se encuentran ante un "
    "dilema. Por un lado, el coste y la complejidad de implantación de las soluciones "
    "empresariales resultan inasumibles. Por otro, gestionar las incidencias mediante "
    "hojas de cálculo, correos electrónicos o sistemas caseros conlleva pérdidas de "
    "información, falta de trazabilidad y una experiencia de usuario deficiente."
)

add_para(
    "Este Trabajo Fin de Ciclo presenta SolveIT, una plataforma SaaS (Software as a "
    "Service) multi-tenant concebida específicamente para cubrir esta necesidad. La idea "
    "de negocio consiste en ofrecer a las PYME un sistema profesional de gestión de "
    "incidencias IT con todas las funcionalidades que un departamento de soporte moderno "
    "requiere, accesible desde el navegador y desde el móvil, con un modelo de pago por "
    "suscripción mensual y sin necesidad de mantenimiento por parte del cliente."
)

add_para(
    "La plataforma propone un servicio en tres capas. La primera capa es la aplicación "
    "móvil multiplataforma, desarrollada con React Native y Expo, que permite a usuarios "
    "y técnicos gestionar incidencias desde cualquier ubicación. La segunda capa es la "
    "aplicación web responsive, construida con Next.js, orientada principalmente a roles "
    "administrativos y a entornos donde el ordenador es la herramienta habitual. La "
    "tercera capa es el backend Node.js con base de datos PostgreSQL, donde reside la "
    "lógica de negocio y se centralizan los datos de todas las empresas cliente bajo una "
    "arquitectura multi-tenant que garantiza el aislamiento total entre organizaciones."
)

add_para(
    "Como elemento diferenciador frente a la competencia, SolveIT incorpora un asistente "
    "de inteligencia artificial basado en el modelo Gemma de Google que asiste a los "
    "usuarios ante problemas técnicos antes de generar un ticket. Este asistente, además "
    "de mejorar la experiencia del empleado al ofrecerle una respuesta inmediata, reduce "
    "la carga sobre el departamento de soporte al filtrar consultas que pueden resolverse "
    "con la documentación existente."
)

add_para(
    "La justificación personal de este proyecto responde a varias motivaciones "
    "interconectadas. Durante la realización de las prácticas en empresa observé de "
    "primera mano cómo los departamentos de IT de las PYME conviven con herramientas "
    "heterogéneas y poco integradas para gestionar el día a día. La idea de construir "
    "una solución que aglutine todas las necesidades en un único producto, accesible y "
    "con tecnologías modernas, surgió de manera natural. A nivel profesional, este "
    "trabajo me permite consolidar las competencias adquiridas en los módulos de "
    "Programación Multimedia y Dispositivos Móviles, Acceso a Datos, Programación de "
    "Servicios y Procesos y Sistemas de Gestión Empresarial, integrándolas en un único "
    "proyecto end-to-end. Además, posiciona mi perfil profesional en un sector con alta "
    "demanda como es el desarrollo full-stack y la integración de inteligencia artificial "
    "en productos reales."
)

add_para(
    "Desde la perspectiva del emprendimiento, SolveIT se presenta como un proyecto "
    "alineado con los Objetivos de Desarrollo Sostenible de la Agenda 2030 (Naciones "
    "Unidas, 2015). En particular, contribuye al ODS 8 (Trabajo decente y crecimiento "
    "económico) al mejorar la productividad de los trabajadores reduciendo el tiempo de "
    "inactividad por incidencias técnicas; al ODS 9 (Industria, innovación e "
    "infraestructura) al democratizar el acceso a herramientas de gestión IT que hasta "
    "ahora estaban reservadas a grandes empresas; y al ODS 12 (Producción y consumo "
    "responsables) al fomentar la digitalización y reducir el uso de papel y "
    "desplazamientos físicos del personal técnico. El proyecto se enmarca, asimismo, "
    "dentro de las economías transformadoras al impulsar la digitalización del tejido "
    "productivo y al ofrecer un modelo de software accesible que reduce la brecha "
    "tecnológica entre grandes corporaciones y PYME."
)

# ─────────────────────────────────────────────────────────────────────────────
# 2. ANÁLISIS
# ─────────────────────────────────────────────────────────────────────────────
add_h1("2. ANÁLISIS")

add_h2("2.1 Estudio de mercado")

add_para(
    "El mercado del IT Service Management (ITSM) ha experimentado un crecimiento sostenido "
    "durante los últimos años. Según el informe de Gartner (2024), el segmento global de "
    "software ITSM superó los 12.000 millones de dólares en 2023 y se proyecta una tasa "
    "anual compuesta de crecimiento (CAGR) cercana al 9% hasta 2028. Este crecimiento "
    "responde a la combinación de tres factores: la digitalización acelerada tras la "
    "pandemia, la adopción de modelos híbridos de trabajo y la creciente complejidad de "
    "las infraestructuras IT."
)

add_para(
    "Para validar la necesidad del producto se ha realizado una investigación cualitativa "
    "basada en observación directa durante el periodo de Formación en Centros de Trabajo "
    "y en conversaciones informales con responsables de soporte de cinco empresas de "
    "entre 15 y 200 empleados. Las conclusiones más relevantes son las siguientes:"
)
add_bullet("Cuatro de las cinco empresas consultadas no disponen de un sistema profesional de tickets. Gestionan las incidencias mediante correo electrónico o aplicaciones de mensajería instantánea.")
add_bullet("El 100% de los responsables consultados reconoce que pierde información sobre incidencias previas, lo que provoca que se repitan diagnósticos para problemas ya resueltos.")
add_bullet("La principal barrera para adoptar una solución profesional es el coste de las licencias y la complejidad de implantación.")
add_bullet("Existe un interés explícito por funcionalidades de inteligencia artificial que reduzcan la carga del equipo de soporte.")

add_h3("Análisis de la competencia")

add_para(
    "ServiceNow es la solución líder en el segmento corporativo. Ofrece una plataforma "
    "completa de gestión de servicios IT, gestión de activos, automatización y analítica. "
    "Su precio parte de aproximadamente 100 dólares por usuario y mes, con costes "
    "adicionales de implantación que pueden superar los 50.000 euros para una empresa "
    "mediana. Su orientación es claramente enterprise y resulta inviable para PYME."
)

add_para(
    "Jira Service Management (Atlassian) es la opción más popular entre empresas "
    "tecnológicas medianas. Su modelo es freemium con un plan gratuito limitado a tres "
    "agentes y planes de pago desde 21 dólares por agente al mes. Aunque su precio es "
    "más accesible que ServiceNow, su curva de aprendizaje es considerable y muchas "
    "funcionalidades requieren plugins de pago."
)

add_para(
    "Freshservice (Freshworks) es una solución cloud orientada a empresas medianas. "
    "Precios desde 19 dólares por agente al mes, interfaz amigable y buena experiencia "
    "móvil. Su limitación principal es que está orientada a empresas con un único tenant "
    "y no incorpora capacidades nativas de IA generativa."
)

add_para(
    "Zammad y osTicket son soluciones open source autoalojadas. El coste de licencia es "
    "nulo pero requieren administración de infraestructura, mantenimiento, parcheado y "
    "conocimiento técnico que la mayoría de PYME no tiene internamente."
)

add_h3("Análisis DAFO de SolveIT")

add_table(
    ["Fortalezas", "Debilidades"],
    [
        ["Tecnología moderna (React Native, Next.js, Node.js)", "Marca desconocida sin reputación previa"],
        ["Asistente de IA integrado de serie", "Equipo inicial reducido (un único desarrollador)"],
        ["Multi-tenant nativo", "Cartera de funcionalidades inferior a ServiceNow"],
        ["Coste reducido para el cliente", "Dependencia de proveedores cloud externos"],
    ]
)

add_table(
    ["Oportunidades", "Amenazas"],
    [
        ["Crecimiento del segmento ITSM", "Entrada de competidores con más recursos"],
        ["Ayudas a la digitalización (Kit Digital)", "Cambios en políticas de precios de cloud o IA"],
        ["Demanda de IA en productos SaaS", "Madurez del open source autoalojable"],
        ["Tejido PYME español infrautilizando ITSM", "Riesgo regulatorio (RGPD, AI Act)"],
    ]
)

add_h2("2.2 Elección de soluciones")

add_para(
    "Tras el análisis de mercado y aplicando la filosofía KISS (Keep It Simple, Stupid), "
    "se ha decidido construir SolveIT como una solución a medida sobre un stack "
    "tecnológico moderno y mayoritariamente abierto, en lugar de adaptar una solución "
    "existente o construir un fork de un sistema open source como Zammad. Las razones de "
    "esta elección son las siguientes."
)

add_para(
    "Re-usabilidad. El stack elegido (Node.js, React Native, Next.js, PostgreSQL) es "
    "ampliamente conocido por la comunidad de desarrolladores, lo que facilita la "
    "incorporación de nuevos miembros al equipo en el futuro y garantiza el soporte a "
    "largo plazo. Los conocimientos adquiridos durante el desarrollo se trasladan sin "
    "fricción a otros proyectos."
)

add_para(
    "Complejidad de instalación. El cliente final no tiene que instalar nada en sus "
    "sistemas. La aplicación móvil se distribuye a través de las tiendas oficiales "
    "(Google Play y App Store) y la aplicación web se accede desde un navegador. El "
    "backend se aloja en infraestructura cloud gestionada por el proveedor del servicio."
)

add_para(
    "Velocidad de desarrollo. React Native y Next.js comparten lenguaje (JavaScript / "
    "TypeScript), lo que permite reutilizar lógica entre web y móvil. Expo simplifica el "
    "ciclo de desarrollo y publicación de la aplicación móvil."
)

add_para(
    "Coste. Todas las tecnologías base son gratuitas y de código abierto. Los servicios "
    "cloud se contratan por uso, lo que permite escalar el coste linealmente con los "
    "ingresos."
)

add_h3("Stack tecnológico definitivo")

add_table(
    ["Capa", "Tecnología", "Motivo"],
    [
        ["Frontend móvil", "React Native + Expo (SDK 54)", "Multiplataforma, rapidez, comunidad"],
        ["Frontend web", "Next.js 14 + TypeScript", "SSR, SEO, ecosistema React"],
        ["Backend", "Node.js + Express", "Curva de aprendizaje suave"],
        ["Base de datos", "PostgreSQL", "ACID, soporte JSON, multi-tenant friendly"],
        ["Hosting", "Railway", "PaaS sencillo, integración Git, escalado"],
        ["IA generativa", "Google Gemma 3 12B", "Calidad alta, plan gratuito viable"],
        ["Email transaccional", "Resend", "API moderna, plan gratuito 100/día"],
        ["Autenticación", "JWT + bcrypt", "Control total, sin dependencia externa"],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 3. VALORACIÓN ECONÓMICA
# ─────────────────────────────────────────────────────────────────────────────
add_h1("3. VALORACIÓN ECONÓMICA")

add_para(
    "En este apartado se desglosan los costes asociados al desarrollo y al mantenimiento "
    "del proyecto durante su primer año de vida. Las cifras corresponden a estimaciones "
    "realistas basadas en los precios actuales del mercado español a mayo de 2026."
)

add_h2("3.1 Hardware")

add_para(
    "El hardware necesario para llevar a cabo el proyecto es mínimo, ya que el stack "
    "elegido permite desarrollar desde un equipo personal sin grandes requisitos."
)

add_table(
    ["Concepto", "Coste (€)", "Vida útil"],
    [
        ["Equipo de desarrollo (portátil 16 GB RAM, SSD 512 GB)", "1.000", "4 años"],
        ["Monitor adicional 24\"", "180", "6 años"],
        ["Smartphone Android para pruebas", "250", "4 años"],
        ["Periféricos (teclado, ratón)", "70", "5 años"],
        ["Total hardware", "1.500", ""],
    ]
)

add_para(
    "Aplicando amortización lineal, el coste imputable al primer año del proyecto es de "
    "aproximadamente 375 €."
)

add_h2("3.2 Licencias y servicios")

add_para(
    "Una de las ventajas de elegir un stack basado en software libre es la práctica "
    "inexistencia de costes de licencia. Los costes recurrentes corresponden "
    "principalmente a servicios cloud y SaaS de terceros."
)

add_table(
    ["Concepto", "Coste mensual (€)", "Coste anual (€)"],
    [
        ["Railway (hosting backend + PostgreSQL)", "5", "60"],
        ["Dominio .com", "–", "12"],
        ["Cuenta Apple Developer", "–", "99"],
        ["Cuenta Google Play Developer (pago único)", "–", "25"],
        ["Resend (plan gratuito)", "0", "0"],
        ["Google Gemini API (plan gratuito)", "0", "0"],
        ["GitHub (repositorio privado)", "0", "0"],
        ["Expo EAS Build (plan gratuito)", "0", "0"],
        ["Total año 1", "", "196"],
    ]
)

add_para(
    "A medida que la plataforma escale, será necesario migrar a planes de pago. Para el "
    "segundo año se estima un coste anual de servicios de aproximadamente 850 €."
)

add_h2("3.3 Personal")

add_para(
    "Durante la fase inicial el desarrollo lo asume el promotor del proyecto. El coste "
    "del personal se valora a precio de mercado para poder construir un análisis de "
    "viabilidad realista, aunque en la práctica se trate de horas autoaplicadas."
)

add_table(
    ["Rol", "Horas", "€/h", "Coste total (€)"],
    [
        ["Desarrollador full-stack", "600", "25", "15.000"],
        ["Diseñador UX/UI (puntual)", "30", "35", "1.050"],
        ["Asesoría legal (constitución, RGPD)", "8", "80", "640"],
        ["Asesoría fiscal (anual)", "–", "–", "720"],
        ["Total personal año 1", "", "", "17.410"],
    ]
)

add_h2("3.4 Costes de mantenimiento")

add_para(
    "Una vez la plataforma está en producción, el mantenimiento incluye actualizaciones "
    "de dependencias, corrección de bugs, atención a incidencias de clientes y "
    "desarrollo evolutivo. Para el primer año se estima una dedicación media de 8 horas "
    "semanales (aproximadamente 384 horas anuales), lo que supone un coste de 9.600 €. "
    "A esto se suma el coste creciente de la infraestructura cloud (estimado en 600 € "
    "adicionales)."
)

add_h3("Resumen económico año 1")

add_table(
    ["Concepto", "Importe (€)"],
    [
        ["Amortización hardware", "375"],
        ["Licencias y servicios", "196"],
        ["Personal desarrollo", "17.410"],
        ["Personal mantenimiento (medio año)", "4.800"],
        ["Total año 1", "22.781"],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 4. VIABILIDAD DEL PROYECTO
# ─────────────────────────────────────────────────────────────────────────────
add_h1("4. VIABILIDAD DEL PROYECTO")

add_para(
    "El análisis de viabilidad se aborda desde tres dimensiones complementarias: "
    "viabilidad técnica, viabilidad económica y viabilidad de mercado."
)

add_h3("Viabilidad técnica")

add_para(
    "La viabilidad técnica está demostrada porque la plataforma se ha desarrollado y "
    "desplegado en producción. La aplicación está accesible públicamente en "
    "solveit-app-production.up.railway.app y la versión móvil es instalable mediante un "
    "APK generado con EAS Build de Expo. El stack elegido es maduro y existe abundante "
    "documentación oficial y comunitaria. La arquitectura multi-tenant basada en "
    "company_id por fila ha sido validada con dos empresas en paralelo (SolveIT Demo y "
    "LKS Next) sin fugas de información entre tenants."
)

add_h3("Viabilidad económica")

add_para(
    "Para evaluar la viabilidad económica se ha construido un modelo de proyección a "
    "tres años con tres planes comerciales:"
)

add_table(
    ["Plan", "Precio (€/mes)", "Usuarios", "Funcionalidades"],
    [
        ["Basic", "19", "Hasta 10", "Tickets, KB, foro"],
        ["Pro", "49", "Hasta 50", "+ IA, dashboard avanzado"],
        ["Enterprise", "149", "Ilimitado", "+ SLA, soporte prioritario"],
    ]
)

add_h3("Proyección a 3 años (escenario realista)")

add_table(
    ["Año", "Clientes", "Ingresos (€)", "Costes (€)", "Resultado (€)"],
    [
        ["1", "8", "4.500", "22.781", "-18.281"],
        ["2", "25", "18.000", "28.000", "-10.000"],
        ["3", "60", "49.000", "38.000", "+11.000"],
    ]
)

add_para(
    "El punto de equilibrio se alcanza durante el tercer año, lo que es coherente con "
    "los plazos habituales en proyectos SaaS B2B (Frasco, 2023). Los dos primeros años "
    "requieren financiación externa o reinversión de horas no remuneradas del promotor."
)

add_h3("Viabilidad de mercado")

add_para(
    "España cuenta con aproximadamente 2,9 millones de PYME activas (Ministerio de "
    "Industria, 2024). Asumiendo conservadoramente que solo el 5% de las PYME con más "
    "de 10 empleados podrían ser clientes potenciales (alrededor de 30.000 empresas), "
    "capturar un 0,2% de ese segmento durante los primeros tres años (60 clientes) es "
    "un objetivo razonable. Los clientes objetivo son PYME españolas del sector "
    "servicios, comercio y manufactura ligera con plantillas entre 10 y 100 empleados."
)

add_h3("Riesgos identificados")

add_table(
    ["Riesgo", "Probabilidad", "Impacto", "Mitigación"],
    [
        ["Fuga de información entre tenants", "Baja", "Alto", "Tests automáticos, auditoría"],
        ["Caída prolongada del proveedor cloud", "Media", "Alto", "Backups diarios, plan migración"],
        ["Cambio en política de precios de IA", "Media", "Medio", "Capa de abstracción multi-proveedor"],
        ["Brecha de seguridad / RGPD", "Baja", "Muy alto", "Cifrado, auditorías, seguro"],
        ["Aparición de competidor con más recursos", "Alta", "Medio", "Foco PYME, precio, IA, soporte"],
        ["Baja adopción inicial", "Media", "Alto", "Freemium, pilotos gratuitos, marketing"],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 5. PUESTA EN MARCHA
# ─────────────────────────────────────────────────────────────────────────────
add_h1("5. PUESTA EN MARCHA")

add_h2("5.1 Forma jurídica y trámites de constitución")

add_para(
    "La forma jurídica seleccionada para la explotación comercial de SolveIT es la "
    "Sociedad Limitada Unipersonal (SLU). Esta elección se justifica por tres motivos. "
    "Primero, limita la responsabilidad del emprendedor al capital aportado, protegiendo "
    "el patrimonio personal frente a posibles reclamaciones. Segundo, transmite mayor "
    "confianza a clientes B2B que otras formas como autónomo. Tercero, facilita la "
    "entrada de socios o inversores en el futuro mediante ampliación de capital."
)

add_para("Los trámites para la constitución son los siguientes:")
add_bullet("Solicitud de denominación social en el Registro Mercantil Central. Coste: 16 €.")
add_bullet("Apertura de cuenta bancaria y depósito del capital social mínimo (3.000 €).")
add_bullet("Otorgamiento de escritura pública ante notario. Coste estimado: 200 €.")
add_bullet("Liquidación del Impuesto de Operaciones Societarias (exento desde 2010).")
add_bullet("Inscripción en el Registro Mercantil. Coste estimado: 100 €.")
add_bullet("Solicitud del NIF definitivo ante la Agencia Tributaria.")
add_bullet("Alta en Hacienda mediante el modelo 036 (Censo de Empresarios).")
add_bullet("Alta en la Seguridad Social (RETA para el administrador).")
add_bullet("Registro como responsable de tratamiento ante la AEPD (RGPD).")

add_h2("5.2 Recursos humanos")

add_para(
    "Durante el primer año la plantilla se compone exclusivamente del promotor en doble "
    "rol de administrador y desarrollador. A partir del segundo año, dependiendo de la "
    "captación de clientes, se planifica la incorporación de un segundo desarrollador "
    "full-stack (junior) y, en el tercer año, de un perfil comercial con conocimientos "
    "técnicos."
)

add_table(
    ["Año", "Plantilla"],
    [
        ["1", "1 desarrollador full-stack (promotor)"],
        ["2", "2 desarrolladores"],
        ["3", "2 desarrolladores + 1 comercial"],
    ]
)

add_para(
    "La organización se diseña para ser plana, con reuniones semanales, uso intensivo "
    "de herramientas de colaboración asíncrona (GitHub, Notion, Slack) y posibilidad de "
    "teletrabajo total para reducir costes inmobiliarios y atraer talento más allá de "
    "la zona geográfica del promotor."
)

add_h2("5.3 Prevención de riesgos laborales")

add_para(
    "Aunque inicialmente la actividad se desarrolle en régimen de autónomo desde "
    "domicilio o espacios de coworking, la incorporación futura de personal asalariado "
    "obliga al cumplimiento de la Ley 31/1995 de Prevención de Riesgos Laborales. Los "
    "principales riesgos identificados en una empresa de desarrollo software son:"
)
add_bullet("Riesgos ergonómicos derivados del trabajo prolongado con pantallas (síndrome del túnel carpiano, dolor lumbar, fatiga visual).")
add_bullet("Riesgo psicosocial por carga mental, plazos ajustados y trabajo aislado en remoto.")
add_bullet("Riesgos eléctricos en puestos de trabajo.")

add_para(
    "Las medidas previstas incluyen la contratación de un Servicio de Prevención Ajeno "
    "que realice la evaluación inicial de riesgos, la formación obligatoria al personal "
    "en uso seguro de pantallas (Real Decreto 488/1997), la dotación de mobiliario "
    "ergonómico y la planificación de pausas activas. En entorno de teletrabajo se "
    "aplicará la Ley 10/2021 de trabajo a distancia, formalizando un acuerdo individual "
    "con cada trabajador."
)

add_h2("5.4 Análisis medioambiental")

add_para(
    "Aunque la actividad de desarrollo de software tiene un impacto medioambiental "
    "directo limitado, sí existen consideraciones relevantes en el ciclo de vida de los "
    "servicios cloud y en los hábitos de trabajo. Se adoptan las siguientes medidas:"
)
add_bullet("Selección de proveedores cloud con compromiso de neutralidad de carbono.")
add_bullet("Optimización del consumo de recursos para reducir el consumo energético del backend.")
add_bullet("Política de cero papel: documentación y facturación digitales.")
add_bullet("Reciclaje del hardware al final de su vida útil.")
add_bullet("Política de teletrabajo que reduce desplazamientos y emisiones de CO₂.")

add_h2("5.5 Plan de financiación")

add_para(
    "La financiación inicial necesaria para superar el periodo de pérdidas (años 1 y 2) "
    "se estima en aproximadamente 30.000 €. El plan combina recursos propios y ajenos:"
)

add_table(
    ["Fuente", "Importe (€)", "Tipo"],
    [
        ["Capital social aportado por promotor", "3.000", "Propios"],
        ["Ahorros del promotor", "5.000", "Propios"],
        ["Préstamo participativo ENISA Jóvenes Emprendedores", "15.000", "Ajenos"],
        ["Subvención Kit Digital (cliente final)", "–", "Indirecta"],
        ["Ingresos por suscripciones año 1-2", "22.500", "Operativos"],
        ["Total financiación cubierta", "45.500", ""],
    ]
)

add_para(
    "Adicionalmente, se valora la solicitud de la Ayuda CDTI Neotec para empresas de "
    "base tecnológica de hasta 250.000 € en caso de validación temprana del modelo. "
    "Los fondos europeos Next Generation vinculados a la transformación digital también "
    "constituyen una vía relevante a explorar a través del programa Kit Digital, que "
    "beneficia indirectamente a SolveIT al subvencionar a sus clientes potenciales "
    "(PYME) en la adopción de soluciones digitales como la nuestra (Red.es, 2024)."
)

# ─────────────────────────────────────────────────────────────────────────────
# 6. EVALUACIÓN FUNCIONAL
# ─────────────────────────────────────────────────────────────────────────────
add_h1("6. EVALUACIÓN FUNCIONAL")

add_para(
    "La plataforma SolveIT ha sido desplegada en producción y se ha validado "
    "funcionalmente sobre dos empresas piloto: SolveIT Demo (interna, 27 usuarios) y "
    "LKS Next (externa, 3 usuarios). El soporte de evaluación es la propia aplicación "
    "corriendo sobre infraestructura Railway. Los resultados de la verificación "
    "funcional se exponen a continuación."
)

add_h3("Funcionalidades verificadas")

add_para(
    "Autenticación y gestión de usuarios. Se ha verificado el inicio de sesión, el "
    "cierre de sesión, la recuperación de contraseña mediante código enviado por correo "
    "electrónico (utilizando Resend como proveedor SMTP-API) y el control de acceso por "
    "roles (superadmin, admin, technician, user). Los mensajes de error se muestran de "
    "forma visible en la pantalla de login."
)

add_para(
    "Gestión de incidencias. Los usuarios pueden crear incidencias indicando título, "
    "descripción, prioridad y categoría. El sistema asigna automáticamente al técnico "
    "con menor carga de trabajo activa. Los técnicos cambian el estado de las "
    "incidencias entre los valores open, in_progress, resolved y closed, y todos los "
    "cambios se registran en el historial."
)

add_para(
    "Sistema de notificaciones. Cuando se crea una incidencia, el técnico asignado "
    "recibe una notificación en la campanita de la aplicación. Cuando el estado cambia, "
    "el creador de la incidencia es notificado. La tabla notifications incluye los "
    "campos title, body y type para construir notificaciones tipadas y enriquecidas."
)

add_para(
    "Base de conocimiento. Los administradores pueden crear, editar y eliminar "
    "artículos. Los usuarios pueden valorar los artículos con una puntuación de 1 a 5 "
    "estrellas y dejar reseñas. Los artículos están vinculados a categorías y soportan "
    "imágenes de portada."
)

add_para(
    "Foro interno. Permite a los empleados publicar mensajes, comentar y reaccionar con "
    "emojis. Es un canal de comunicación interno alternativo al correo electrónico para "
    "temas no estrictamente vinculados a una incidencia."
)

add_para(
    "Asistente de inteligencia artificial. Se ha integrado el modelo Gemma 3 12B "
    "mediante la API gratuita de Google Generative Language. El asistente recibe la "
    "pregunta del usuario y, opcionalmente, contexto de la base de conocimiento, y "
    "devuelve una respuesta detallada en español. El acceso al asistente se realiza "
    "mediante un botón flotante (FAB) presente en todas las pantallas principales."
)

add_para(
    "Arquitectura multi-tenant. Cada fila de las tablas incidents, users, "
    "knowledge_base, forum_posts y notifications incluye un campo company_id que se "
    "filtra automáticamente en cada consulta a través del middleware tenant.js. Se ha "
    "verificado que un administrador de la empresa A no puede ver datos de la empresa B "
    "aunque manipule el token JWT."
)

add_para(
    "Panel de SuperAdmin. Permite al desarrollador (rol superadmin) crear nuevas "
    "empresas cliente, asignar el plan contratado (Basic, Pro o Enterprise), activar o "
    "desactivar empresas y crear el usuario administrador inicial."
)

add_h3("Métricas y herramientas de validación")

add_bullet("Pruebas manuales sobre dispositivos Android reales (Xiaomi y Samsung) y sobre Expo Web en navegador Chrome.")
add_bullet("Logs de Railway para verificar el comportamiento del backend ante peticiones reales.")
add_bullet("Network tab del navegador para verificar tiempos de respuesta y códigos HTTP.")
add_bullet("Validación de aislamiento multi-tenant mediante usuarios de prueba en distintas empresas.")
add_bullet("Inspección de la base de datos vía cliente PostgreSQL para confirmar persistencia con company_id correcto.")

add_h3("Consecución de objetivos")

add_table(
    ["Objetivo planteado", "Resultado"],
    [
        ["Aplicación móvil funcional en Android", "✓ APK generado y probado"],
        ["Aplicación web responsive", "✓ Desplegada en Railway"],
        ["Backend API REST", "✓ 11 grupos de rutas operativos"],
        ["Multi-tenant con aislamiento real", "✓ Verificado con dos empresas"],
        ["Asistente IA integrado", "✓ Gemma 3 12B funcionando"],
        ["Recuperación de contraseña", "✓ Resend operativo"],
        ["Sistema de notificaciones", "✓ Funcionando en tiempo real"],
        ["Documentación visual de soporte", "✓ Capturas y vídeo demostrativo"],
    ]
)

add_para(
    "Todos los objetivos planteados al inicio del proyecto se han alcanzado "
    "satisfactoriamente, lo que permite afirmar que la plataforma es funcionalmente "
    "válida para su comercialización a clientes piloto."
)

# ─────────────────────────────────────────────────────────────────────────────
# 7. CONCLUSIÓN
# ─────────────────────────────────────────────────────────────────────────────
add_h1("7. CONCLUSIÓN")

add_para(
    "El presente trabajo ha consistido en el diseño, desarrollo, despliegue y "
    "validación de SolveIT, una plataforma SaaS multi-tenant para la gestión de "
    "incidencias informáticas en pequeñas y medianas empresas. A lo largo del proyecto "
    "se ha cubierto el ciclo completo de un producto software, desde la identificación "
    "de la necesidad de mercado hasta la puesta en producción real con clientes piloto."
)

add_para(
    "Las principales aportaciones del trabajo son tres. En primer lugar, se demuestra "
    "que es viable construir un producto SaaS completo y profesional con un stack "
    "mayoritariamente abierto y servicios cloud gratuitos o de bajo coste, abriendo la "
    "puerta a que otros emprendedores con recursos limitados aborden proyectos de "
    "complejidad similar. En segundo lugar, se valida que la integración de "
    "inteligencia artificial generativa en productos B2B no es solo un reclamo de "
    "marketing sino una funcionalidad que aporta valor real reduciendo la carga del "
    "equipo de soporte. En tercer lugar, se confirma que el modelo multi-tenant basado "
    "en discriminador por columna es suficiente para los volúmenes esperados en el "
    "segmento PYME, sin necesidad de bases de datos separadas por cliente."
)

add_para(
    "A nivel personal y profesional, el trabajo ha permitido consolidar las competencias "
    "adquiridas durante el ciclo formativo, especialmente en programación de servicios "
    "y procesos, acceso a datos, programación multimedia y dispositivos móviles. La "
    "integración end-to-end de tecnologías heterogéneas (backend, frontend web, móvil, "
    "IA, cloud) refleja el perfil profesional demandado actualmente en el mercado "
    "laboral."
)

add_para(
    "El proyecto está alineado con la Agenda 2030 y los Objetivos de Desarrollo "
    "Sostenible, especialmente el ODS 8, el ODS 9 y el ODS 12, al contribuir a la "
    "digitalización del tejido empresarial español, mejorar la productividad de los "
    "trabajadores y democratizar el acceso a herramientas IT que tradicionalmente solo "
    "estaban al alcance de grandes corporaciones. Se sitúa, asimismo, en el ámbito de "
    "las economías transformadoras al ofrecer un producto accesible que reduce la "
    "brecha tecnológica entre PYME y grandes empresas."
)

add_para(
    "Como líneas futuras de trabajo se proponen las siguientes: la incorporación de un "
    "módulo de gestión de activos (CMDB) que permita asociar incidencias con "
    "dispositivos y software concretos; la integración con Microsoft Teams y Slack "
    "para crear y consultar tickets desde el chat corporativo; el desarrollo de una API "
    "pública que permita a clientes con necesidades específicas integrar SolveIT con "
    "sus sistemas internos; y la incorporación de un módulo de SLA (Service Level "
    "Agreement) para empresas que necesiten garantías formales de tiempos de respuesta."
)

add_para(
    "Para abrir el debate con el tribunal se plantean las siguientes preguntas. "
    "¿Hasta qué punto el uso de modelos de IA gratuitos pero externos compromete la "
    "sostenibilidad a largo plazo de un producto SaaS? ¿Es ética la asignación "
    "automática de incidencias al técnico con menor carga si esto puede generar "
    "desigualdades en la valoración del rendimiento de los empleados? ¿Cómo debe "
    "SolveIT evolucionar para cumplir con los requisitos del próximo Reglamento Europeo "
    "de Inteligencia Artificial?"
)

# ─────────────────────────────────────────────────────────────────────────────
# 8. BIBLIOGRAFÍA
# ─────────────────────────────────────────────────────────────────────────────
add_h1("8. BIBLIOGRAFÍA")

referencias = [
    "Atlassian. (2024). Jira Service Management documentation. Atlassian Documentation. https://www.atlassian.com/software/jira/service-management",
    "Frasco, S. (2023). SaaS metrics 2.0: A guide to measuring and improving what matters. For Entrepreneurs. https://www.forentrepreneurs.com/saas-metrics-2/",
    "Gartner. (2024). Magic Quadrant for IT Service Management Platforms. Gartner Research.",
    "Google. (2024). Gemma: Open models from Google. Google for Developers. https://ai.google.dev/gemma",
    "Instituto Nacional de Estadística. (2024). Encuesta sobre el uso de Tecnologías de la Información y las Comunicaciones (TIC) y del comercio electrónico en las empresas 2023-2024. INE. https://www.ine.es/",
    "Ley 10/2021, de 9 de julio, de trabajo a distancia. Boletín Oficial del Estado, 164, de 10 de julio de 2021.",
    "Ley 31/1995, de 8 de noviembre, de Prevención de Riesgos Laborales. Boletín Oficial del Estado, 269, de 10 de noviembre de 1995.",
    "Ley Orgánica 3/2018, de 5 de diciembre, de Protección de Datos Personales y garantía de los derechos digitales. Boletín Oficial del Estado, 294, de 6 de diciembre de 2018.",
    "Ministerio de Industria, Comercio y Turismo. (2024). Cifras PYME: datos enero 2024. Dirección General de Industria y de la Pequeña y Mediana Empresa. https://industria.gob.es/",
    "Naciones Unidas. (2015). Transformar nuestro mundo: la Agenda 2030 para el Desarrollo Sostenible. Resolución A/RES/70/1.",
    "Real Decreto 488/1997, de 14 de abril, sobre disposiciones mínimas de seguridad y salud relativas al trabajo con equipos que incluyen pantallas de visualización. Boletín Oficial del Estado, 97, de 23 de abril de 1997.",
    "Reglamento (UE) 2016/679 del Parlamento Europeo y del Consejo, de 27 de abril de 2016, relativo a la protección de las personas físicas en lo que respecta al tratamiento de datos personales (RGPD). Diario Oficial de la Unión Europea, L 119, de 4 de mayo de 2016.",
    "Red.es. (2024). Programa Kit Digital: bases reguladoras y convocatorias. Ministerio para la Transformación Digital. https://www.acelerapyme.gob.es/kit-digital",
    "ServiceNow. (2024). ServiceNow IT Service Management. ServiceNow. https://www.servicenow.com/products/itsm.html",
]
for ref in referencias:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(ref)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# 9. ANEXOS
# ─────────────────────────────────────────────────────────────────────────────
add_h1("9. ANEXOS")

add_h2("Anexo I. Capturas de pantalla de la aplicación")
add_para(
    "[INSERTAR AQUÍ las capturas de las pantallas principales: login, lista de "
    "incidencias, detalle de incidencia, base de conocimiento, foro, asistente IA, "
    "panel SuperAdmin]",
    italic=True
)

add_h2("Anexo II. Diagrama de la arquitectura")
add_para(
    "[INSERTAR AQUÍ el diagrama mostrando: app móvil → API REST (Railway) → PostgreSQL, "
    "con conexiones externas a Resend (email) y Google Gemini (IA)]",
    italic=True
)

add_h2("Anexo III. Modelo entidad-relación de la base de datos")
add_para(
    "[INSERTAR AQUÍ el diagrama ER con las tablas: companies, users, incidents, "
    "comments, incident_history, notifications, knowledge_base, knowledge_ratings, "
    "forum_posts, forum_comments, forum_reactions, password_resets, ai_conversations]",
    italic=True
)

add_h2("Anexo IV. Repositorio del código fuente")
add_para("https://github.com/r-aguado/solveit-app")

add_h2("Anexo V. URL de la aplicación en producción")
add_bullet("API backend: https://solveit-app-production.up.railway.app")
add_bullet("Aplicación web: [URL de despliegue de Next.js]")
add_bullet("APK Android: enlace generado por EAS Build de Expo")

add_h2("Anexo VI. Credenciales de prueba")
add_table(
    ["Rol", "Email", "Contraseña"],
    [
        ["SuperAdmin", "superadmin@solveit.com", "superadmin1234"],
        ["Admin (Demo)", "admin@solveit.com", "admin1234"],
        ["Usuarios sembrados", "*@solveit.com", "password"],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Guardar
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_DOCX)
print(f"OK Word generado: {OUTPUT_DOCX}")
